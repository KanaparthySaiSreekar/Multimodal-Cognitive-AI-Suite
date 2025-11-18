"""Data ingestion utilities for loading various file formats."""

import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import pandas as pd
import PyPDF2
import pdfplumber
from docx import Document
from pdf2image import convert_from_path
from PIL import Image

from ..utils.logger import get_logger

logger = get_logger(__name__)


class DataIngestion:
    """Handle data ingestion from various sources and formats."""

    def __init__(self, data_dir: Optional[Union[str, Path]] = None):
        """
        Initialize data ingestion.

        Args:
            data_dir: Base directory for data files
        """
        self.data_dir = Path(data_dir) if data_dir else Path.cwd()

    def load_pdf(self, file_path: Union[str, Path], method: str = "pdfplumber") -> Dict[str, any]:
        """
        Load PDF file.

        Args:
            file_path: Path to PDF file
            method: Extraction method ('pdfplumber' or 'pypdf2')

        Returns:
            Dictionary with text and metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        logger.info(f"Loading PDF: {file_path}")

        if method == "pdfplumber":
            return self._load_pdf_pdfplumber(file_path)
        elif method == "pypdf2":
            return self._load_pdf_pypdf2(file_path)
        else:
            raise ValueError(f"Unknown PDF extraction method: {method}")

    def _load_pdf_pdfplumber(self, file_path: Path) -> Dict[str, any]:
        """Load PDF using pdfplumber."""
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            return {
                "text": text.strip(),
                "num_pages": len(pdf.pages),
                "metadata": pdf.metadata,
                "file_path": str(file_path),
            }

    def _load_pdf_pypdf2(self, file_path: Path) -> Dict[str, any]:
        """Load PDF using PyPDF2."""
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            return {
                "text": text.strip(),
                "num_pages": len(pdf_reader.pages),
                "metadata": pdf_reader.metadata,
                "file_path": str(file_path),
            }

    def pdf_to_images(
        self, file_path: Union[str, Path], dpi: int = 200
    ) -> List[Image.Image]:
        """
        Convert PDF pages to images.

        Args:
            file_path: Path to PDF file
            dpi: DPI for image conversion

        Returns:
            List of PIL Images
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        logger.info(f"Converting PDF to images: {file_path}")
        images = convert_from_path(file_path, dpi=dpi)

        return images

    def load_docx(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """
        Load DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with text and metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        logger.info(f"Loading DOCX: {file_path}")

        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        return {
            "text": text.strip(),
            "num_paragraphs": len(doc.paragraphs),
            "file_path": str(file_path),
        }

    def load_text_file(self, file_path: Union[str, Path], encoding: str = "utf-8") -> str:
        """
        Load plain text file.

        Args:
            file_path: Path to text file
            encoding: File encoding

        Returns:
            Text content
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Text file not found: {file_path}")

        logger.info(f"Loading text file: {file_path}")

        with open(file_path, "r", encoding=encoding) as f:
            return f.read().strip()

    def load_image(self, file_path: Union[str, Path]) -> Image.Image:
        """
        Load image file.

        Args:
            file_path: Path to image file

        Returns:
            PIL Image
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Image file not found: {file_path}")

        logger.info(f"Loading image: {file_path}")
        return Image.open(file_path).convert("RGB")

    def load_csv(
        self, file_path: Union[str, Path], text_column: str, label_column: Optional[str] = None
    ) -> pd.DataFrame:
        """
        Load CSV file with text data.

        Args:
            file_path: Path to CSV file
            text_column: Name of text column
            label_column: Name of label column (optional)

        Returns:
            DataFrame with text and labels
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"CSV file not found: {file_path}")

        logger.info(f"Loading CSV: {file_path}")

        df = pd.read_csv(file_path)

        if text_column not in df.columns:
            raise ValueError(f"Text column '{text_column}' not found in CSV")

        if label_column and label_column not in df.columns:
            raise ValueError(f"Label column '{label_column}' not found in CSV")

        return df

    def load_json(self, file_path: Union[str, Path]) -> Union[Dict, List]:
        """
        Load JSON file.

        Args:
            file_path: Path to JSON file

        Returns:
            Parsed JSON data
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"JSON file not found: {file_path}")

        logger.info(f"Loading JSON: {file_path}")

        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def load_dataset(
        self, data_path: Union[str, Path], file_type: Optional[str] = None
    ) -> Union[Dict, pd.DataFrame, List]:
        """
        Load dataset based on file type.

        Args:
            data_path: Path to data file or directory
            file_type: File type override

        Returns:
            Loaded data
        """
        data_path = Path(data_path)

        if not data_path.exists():
            raise FileNotFoundError(f"Data path not found: {data_path}")

        # Determine file type
        if file_type is None:
            file_type = data_path.suffix.lower()

        # Load based on file type
        if file_type in [".csv"]:
            return pd.read_csv(data_path)
        elif file_type in [".json", ".jsonl"]:
            return self.load_json(data_path)
        elif file_type in [".pdf"]:
            return self.load_pdf(data_path)
        elif file_type in [".docx"]:
            return self.load_docx(data_path)
        elif file_type in [".txt"]:
            return self.load_text_file(data_path)
        elif file_type in [".jpg", ".jpeg", ".png", ".bmp"]:
            return self.load_image(data_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def batch_load_images(self, directory: Union[str, Path], extensions: Optional[List[str]] = None) -> List[Tuple[str, Image.Image]]:
        """
        Load all images from a directory.

        Args:
            directory: Directory containing images
            extensions: List of image extensions to load

        Returns:
            List of tuples (file_path, image)
        """
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        if extensions is None:
            extensions = [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]

        images = []
        for ext in extensions:
            for file_path in directory.glob(f"*{ext}"):
                try:
                    image = self.load_image(file_path)
                    images.append((str(file_path), image))
                except Exception as e:
                    logger.error(f"Error loading image {file_path}: {e}")

        logger.info(f"Loaded {len(images)} images from {directory}")
        return images

    def batch_load_documents(
        self, directory: Union[str, Path], extensions: Optional[List[str]] = None
    ) -> List[Dict[str, any]]:
        """
        Load all documents from a directory.

        Args:
            directory: Directory containing documents
            extensions: List of document extensions to load

        Returns:
            List of document dictionaries
        """
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        if extensions is None:
            extensions = [".pdf", ".docx", ".txt"]

        documents = []
        for ext in extensions:
            for file_path in directory.glob(f"*{ext}"):
                try:
                    doc_data = self.load_dataset(file_path)
                    if isinstance(doc_data, dict):
                        documents.append(doc_data)
                    else:
                        documents.append({"text": str(doc_data), "file_path": str(file_path)})
                except Exception as e:
                    logger.error(f"Error loading document {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents
